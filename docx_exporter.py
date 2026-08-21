import io
import re
import base64
from datetime import datetime
from typing import Callable, Optional

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from docx.opc.part import Part
from docx.opc.packuri import PackURI
from PIL import Image


def _set_cell_background(cell, hex_color: str):
    """Set background color of a table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)


def _set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set inner margins (padding) for a table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def _create_fallback_png() -> bytes:
    """Create a 1x1 transparent PNG for Word OpenXML fallback."""
    img = Image.new('RGBA', (1, 1), (255, 255, 255, 0))
    buf = io.BytesIO()
    img.save(buf, format='PNG')
    return buf.getvalue()


def _extract_svg_dimensions(svg_str: str, max_width_in=5.5, max_height_in=6.0):
    """Calculate width and height in inches from SVG attributes or viewBox."""
    w, h = 600.0, 350.0
    vb_m = re.search(r'viewBox=["\']([^"\']+)["\']', svg_str, re.IGNORECASE)
    if vb_m:
        parts = [float(x) for x in re.split(r'[\s,]+', vb_m.group(1).strip()) if x]
        if len(parts) >= 4 and parts[2] > 0 and parts[3] > 0:
            w, h = parts[2], parts[3]
    else:
        w_m = re.search(r'width=["\']([0-9.]+)(?:px)?["\']', svg_str, re.IGNORECASE)
        h_m = re.search(r'height=["\']([0-9.]+)(?:px)?["\']', svg_str, re.IGNORECASE)
        if w_m and h_m:
            try:
                w, h = float(w_m.group(1)), float(h_m.group(1))
            except ValueError:
                pass

    aspect = h / max(w, 1.0)
    width_in = min(max_width_in, 5.5)
    height_in = width_in * aspect
    if height_in > max_height_in:
        height_in = max_height_in
        width_in = height_in / aspect
    return width_in, height_in


def _svg_to_png_bytes(svg_str: str) -> Optional[bytes]:
    """Convert SVG string to high-quality PNG bytes using svglib/reportlab or cairosvg."""
    if not svg_str or not svg_str.strip().startswith("<svg"):
        return None
    try:
        from svglib.svglib import svg2rlg
        from reportlab.graphics import renderPM
        drawing = svg2rlg(io.StringIO(svg_str))
        if drawing:
            png_bytes = renderPM.drawToString(drawing, fmt='PNG')
            if png_bytes and len(png_bytes) > 50:
                return png_bytes
    except Exception:
        pass

    try:
        import cairosvg
        png_bytes = cairosvg.svg2png(bytestring=svg_str.encode('utf-8'))
        if png_bytes and len(png_bytes) > 50:
            return png_bytes
    except Exception:
        pass

    return None


def _add_svg_to_doc(doc, svg_str: str, title="Drawing"):
    """Embed an SVG drawing into a python-docx Document as high-fidelity PNG and/or native SVG."""
    try:
        # 1. Rasterize to PNG first for universal Microsoft Word / LibreOffice compatibility
        png_bytes = _svg_to_png_bytes(svg_str)
        if png_bytes:
            return _add_raster_image_to_doc(doc, png_bytes, alt_text=f"🎨 {title}" if title else "")

        # 2. Fallback to OpenXML SVG markup
        width_in, height_in = _extract_svg_dimensions(svg_str)
        svg_bytes = svg_str.encode('utf-8') if isinstance(svg_str, str) else svg_str
        fallback_png_bytes = _create_fallback_png()

        doc_part = doc.part
        package = doc_part.package

        # 1. Add SVG image part
        svg_num = len([p for p in package.parts if 'image' in p.partname]) + 1
        svg_partname = PackURI(f'/word/media/image_svg_{svg_num}.svg')
        svg_part = Part(svg_partname, 'image/svg+xml', svg_bytes, package)
        package.parts.append(svg_part)
        svg_rId = doc_part.relate_to(svg_part, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image')

        # 2. Add fallback PNG image part
        png_partname = PackURI(f'/word/media/image_png_{svg_num}.png')
        png_part = Part(png_partname, 'image/png', fallback_png_bytes, package)
        package.parts.append(png_part)
        png_rId = doc_part.relate_to(png_part, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image')

        # Dimensions in EMUs (1 inch = 914400 EMUs)
        cx = int(width_in * 914400)
        cy = int(height_in * 914400)

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)

        drawing_xml = f'''
        <w:drawing {nsdecls("w")}>
          <wp:inline distT="0" distB="0" distL="0" distR="0" {nsdecls("wp")}>
            <wp:extent cx="{cx}" cy="{cy}"/>
            <wp:docPr id="{svg_num + 300}" name="{title}"/>
            <a:graphic {nsdecls("a")}>
              <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
                <pic:pic {nsdecls("pic")}>
                  <pic:nvPicPr>
                    <pic:cNvPr id="0" name="{title}"/>
                    <pic:cNvPicPr/>
                  </pic:nvPicPr>
                  <pic:blipFill>
                    <a:blip r:embed="{png_rId}" {nsdecls("r")}>
                      <a:extLst>
                        <a:ext uri="{{96DAC542-7B16-430E-8263-3401B00B12A0}}">
                          <asvg:svgBlip r:embed="{svg_rId}" xmlns:asvg="http://schemas.microsoft.com/office/drawing/2016/SVG/main"/>
                        </a:ext>
                      </a:extLst>
                    </a:blip>
                    <a:stretch>
                      <a:fillRect/>
                    </a:stretch>
                  </pic:blipFill>
                  <pic:spPr>
                    <a:xfrm>
                      <a:off x="0" y="0"/>
                      <a:ext cx="{cx}" cy="{cy}"/>
                    </a:xfrm>
                    <a:prstGeom prst="rect">
                      <a:avLst/>
                    </a:prstGeom>
                  </pic:spPr>
                </pic:pic>
              </a:graphicData>
            </a:graphic>
          </wp:inline>
        </w:drawing>
        '''
        run = p.add_run()
        run._r.append(parse_xml(drawing_xml))

        if title and title.strip():
            p_cap = doc.add_paragraph()
            p_cap.paragraph_format.space_before = Pt(0)
            p_cap.paragraph_format.space_after = Pt(8)
            r_cap = p_cap.add_run(f"🎨 {title.strip()}")
            r_cap.font.size = Pt(9)
            r_cap.font.color.rgb = RGBColor(100, 116, 139)
            r_cap.italic = True
        return p
    except Exception as e:
        p_err = doc.add_paragraph()
        r_err = p_err.add_run(f"🎨 [Gambar/Canvas: {title}]")
        r_err.font.color.rgb = RGBColor(100, 116, 139)
        return p_err


def _add_raster_image_to_doc(doc, img_bytes: bytes, alt_text="", max_width_in=5.5, max_height_in=6.0):
    """Add a raster image (PNG, JPEG, WebP, etc.) to Word document."""
    try:
        pil_img = Image.open(io.BytesIO(img_bytes))
        w, h = pil_img.size
        fmt = (pil_img.format or 'PNG').upper()
        if fmt not in ('PNG', 'JPEG', 'JPG', 'BMP', 'GIF'):
            # Convert WebP or other unsupported formats to PNG in memory
            out_buf = io.BytesIO()
            pil_img.save(out_buf, format='PNG')
            img_bytes = out_buf.getvalue()

        aspect = h / max(w, 1.0)
        width_in = min(max_width_in, 5.5)
        height_in = width_in * aspect
        if height_in > max_height_in:
            height_in = max_height_in
            width_in = height_in / aspect

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run()
        run.add_picture(io.BytesIO(img_bytes), width=Inches(width_in), height=Inches(height_in))

        alt_clean = alt_text.strip() if alt_text else ""
        is_raw_fn = alt_clean.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.webp', '.bmp', '.svg', '.ico')) or alt_clean.lower() in ('image', 'screenshot', 'photo', 'foto', 'img')
        if alt_clean and not alt_clean.startswith("http") and not alt_clean.startswith("/") and not is_raw_fn:
            p_cap = doc.add_paragraph()
            p_cap.paragraph_format.space_before = Pt(0)
            p_cap.paragraph_format.space_after = Pt(8)
            r_cap = p_cap.add_run(alt_clean)
            r_cap.font.size = Pt(9)
            r_cap.font.color.rgb = RGBColor(100, 116, 139)
            r_cap.italic = True
        return p
    except Exception as e:
        if alt_text and alt_text.strip():
            p_alt = doc.add_paragraph()
            r_alt = p_alt.add_run(f"🖼️ [{alt_text.strip()}]")
            r_alt.font.color.rgb = RGBColor(100, 116, 139)
            return p_alt
        return None


def _resolve_image_bytes(src: str, image_resolver: Optional[Callable[..., Optional[bytes]]] = None, alt: str = None) -> Optional[bytes]:
    """Resolve image source (base64 data URI, resolver callback, or HTTP URL) into bytes."""
    if not src:
        return None

    src = src.replace('\\', '').strip()
    # 1. Base64 Data URI
    if src.startswith("data:image/") and ";base64," in src:
        try:
            _, b64_data = src.split(";base64,", 1)
            return base64.b64decode(b64_data)
        except Exception:
            return None

    # 2. Custom Resolver callback (e.g. Nextcloud attachments / local DB / filename match)
    if image_resolver:
        try:
            resolved = image_resolver(src, alt=alt)
            if resolved:
                return resolved
        except TypeError:
            try:
                resolved = image_resolver(src)
                if resolved:
                    return resolved
            except Exception:
                pass
        except Exception:
            pass

    # 3. HTTP / HTTPS external URL
    if src.startswith("http://") or src.startswith("https://"):
        try:
            import requests
            r = requests.get(src, timeout=8)
            if r.status_code == 200:
                return r.content
        except Exception:
            pass

def _parse_standalone_image(line: str) -> Optional[tuple[str, str]]:
    """Detect if line is a standalone image reference. Returns (src, alt) or None."""
    line = line.strip()
    if not line:
        return None

    # 1. HTML <img> Tag: <img src="..." alt="..." />
    m_html = re.search(r'<img\s+[^>]*src\s*=\s*(?:"([^"]*)"|\'([^\']*)\')[^>]*>', line, re.IGNORECASE)
    if m_html:
        src = m_html.group(1) or m_html.group(2) or ""
        alt_m = re.search(r'alt\s*=\s*(?:"([^"]*)"|\'([^\']*)\')', line, re.IGNORECASE)
        alt = (alt_m.group(1) or alt_m.group(2)) if alt_m else ""
        return src.strip(), alt.strip()

    # 2. Markdown Image with URL: ![alt](src) or \!\[alt\]\(src\)
    m_md = re.match(r'^\s*\\?!\s*\\?\[\s*([^\]]*)\s*\\?\]\s*\\?\(\s*([^)]*)\s*\\?\)\s*$', line)
    if m_md:
        return m_md.group(2).strip(), m_md.group(1).strip()

    # 3. Standalone image link with URL: [image.png](src)
    m_link_img = re.match(r'^\s*\\?\[\s*([^\]]+\.(?:png|jpg|jpeg|gif|webp|bmp|svg|ico))\s*\\?\]\s*\\?\(\s*([^)]*)\s*\\?\)\s*$', line, re.IGNORECASE)
    if m_link_img:
        return m_link_img.group(2).strip(), m_link_img.group(1).strip()

    # 4. Standalone !image.png or \!image.png
    m_bang_img = re.match(r'^\s*\\?!([^\s\[\]\(\)]+\.(?:png|jpg|jpeg|gif|webp|bmp|svg|ico))\s*$', line, re.IGNORECASE)
    if m_bang_img:
        fn = m_bang_img.group(1).strip()
        return fn, fn

    # 5. Standalone ![image.png] without URL
    m_bracket_bang = re.match(r'^\s*\\?!\s*\\?\[\s*([^\]]+)\s*\\?\]\s*$', line, re.IGNORECASE)
    if m_bracket_bang:
        fn = m_bracket_bang.group(1).strip()
        return fn, fn

    # 6. Standalone [image.png] without URL
    m_bracket_img = re.match(r'^\s*\\?\[\s*([^\]]+\.(?:png|jpg|jpeg|gif|webp|bmp|svg|ico))\s*\\?\]\s*$', line, re.IGNORECASE)
    if m_bracket_img:
        fn = m_bracket_img.group(1).strip()
        return fn, fn

    # 7. Standalone image filename on its own line (e.g. image.png)
    m_raw_img = re.match(r'^\s*([^\s\[\]\(\)]+\.(?:png|jpg|jpeg|gif|webp|bmp|svg|ico))\s*$', line, re.IGNORECASE)
    if m_raw_img:
        fn = m_raw_img.group(1).strip()
        return fn, fn

    return None


def _clean_markdown_text(text: str) -> str:
    """Preprocess markdown text: normalize line breaks, clean HTML tags, and normalize escapes."""
    if not text:
        return ""
    # Replace <br>, <br/>, <br /> with newlines
    text = re.sub(r'<br\s*/?>', '\n', text, flags=re.IGNORECASE)
    # Remove excessive blank lines (> 2 in a row)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text


def _add_styled_runs(paragraph, text: str):
    """Parse inline markdown (bold, italic, strikethrough, code, links) and add formatted runs."""
    token_pattern = re.compile(
        r'(`[^`]+`)'
        r'|(\*\*[^*]+\*\*|__[^_]+__)'
        r'|(\*[^*]+\*|_[^_]+_)'
        r'|(~~[^~]+~~)'
        r'|(\[[^\]]+\]\([^)]+\))'
        r'|(\\[\[\]\(\)*_~`#])'
    )

    def _unescape_plain(s: str) -> str:
        return re.sub(r'\\([\[\]\(\)*_~`#])', r'\1', s)

    last_idx = 0
    for m in token_pattern.finditer(text):
        start, end = m.span()
        if start > last_idx:
            paragraph.add_run(_unescape_plain(text[last_idx:start]))

        token = m.group(0)
        if token.startswith('\\') and len(token) == 2:
            # Escaped character
            paragraph.add_run(token[1])
        elif token.startswith('`') and token.endswith('`'):
            # Inline code
            r = paragraph.add_run(token[1:-1])
            r.font.name = 'Consolas'
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(180, 40, 40)
        elif (token.startswith('**') and token.endswith('**')) or (token.startswith('__') and token.endswith('__')):
            # Bold
            r = paragraph.add_run(_unescape_plain(token[2:-2]))
            r.bold = True
        elif (token.startswith('*') and token.endswith('*')) or (token.startswith('_') and token.endswith('_')):
            # Italic
            r = paragraph.add_run(_unescape_plain(token[1:-1]))
            r.italic = True
        elif token.startswith('~~') and token.endswith('~~'):
            # Strikethrough
            r = paragraph.add_run(_unescape_plain(token[2:-2]))
            r.font.strike = True
        elif token.startswith('[') and '](' in token and token.endswith(')'):
            # Link [text](url)
            link_m = re.match(r'\[([^\]]+)\]\(([^)]+)\)', token)
            if link_m:
                link_text, _ = link_m.groups()
                r = paragraph.add_run(_unescape_plain(link_text))
                r.font.color.rgb = RGBColor(37, 99, 235)  # Blue
                r.underline = True
            else:
                paragraph.add_run(_unescape_plain(token))

        last_idx = end

    if last_idx < len(text):
        paragraph.add_run(_unescape_plain(text[last_idx:]))


def markdown_to_docx(
    title: str,
    content: str,
    meta: dict = None,
    drawing_resolver: Optional[Callable[..., Optional[dict]]] = None,
    image_resolver: Optional[Callable[[str], Optional[bytes]]] = None
) -> io.BytesIO:
    """Convert Markdown content and metadata into a formatted Word (.docx) document with images and drawings."""
    doc = docx.Document()

    # Configure default document styles
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Segoe UI'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(30, 41, 59)

    # Set 1 inch page margins
    for s in doc.sections:
        s.top_margin = Inches(1.0)
        s.bottom_margin = Inches(1.0)
        s.left_margin = Inches(1.0)
        s.right_margin = Inches(1.0)

    # 1. Document Title
    doc_title = title.strip() if title else "Catatan TaskFlow"
    p_title = doc.add_paragraph()
    r_title = p_title.add_run(doc_title)
    r_title.font.name = 'Segoe UI Semibold'
    r_title.font.size = Pt(22)
    r_title.font.color.rgb = RGBColor(15, 23, 42)
    p_title.paragraph_format.space_after = Pt(4)
    p_title.paragraph_format.space_before = Pt(0)

    # 2. Metadata Subtitle (Tags, Date)
    meta = meta or {}
    meta_parts = []
    if meta.get("updated_at"):
        try:
            dt = datetime.fromisoformat(meta["updated_at"].replace("Z", "+00:00"))
            meta_parts.append("Terakhir diperbarui: " + dt.strftime("%d %b %Y, %H:%M"))
        except Exception:
            meta_parts.append("Tanggal: " + str(meta["updated_at"]))

    if meta.get("tags") and isinstance(meta["tags"], list):
        tags_str = "  ".join(f"#{t}" for t in meta["tags"] if t)
        if tags_str:
            meta_parts.append(tags_str)

    if meta_parts:
        p_meta = doc.add_paragraph()
        r_meta = p_meta.add_run(" • ".join(meta_parts))
        r_meta.font.size = Pt(9.5)
        r_meta.font.color.rgb = RGBColor(100, 116, 139)
        p_meta.paragraph_format.space_after = Pt(16)

    # Add divider line
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_after = Pt(14)
    p_div_border = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="E2E8F0"/></w:pBdr>')
    p_div._element.get_or_add_pPr().append(p_div_border)

    # 3. Clean and normalize content
    cleaned_content = _clean_markdown_text(content or "")
    lines = cleaned_content.split('\n')
    i = 0
    n = len(lines)

    while i < n:
        raw_line = lines[i]
        line = raw_line.strip()

        # Blank line
        if not line:
            i += 1
            continue

        # Code Block: ```
        if line.startswith('```'):
            code_lines = []
            i += 1
            while i < n and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # Skip closing ```

            code_text = '\n'.join(code_lines)
            p_code = doc.add_paragraph()
            p_code.paragraph_format.left_indent = Inches(0.2)
            p_code.paragraph_format.right_indent = Inches(0.2)
            p_code.paragraph_format.space_before = Pt(6)
            p_code.paragraph_format.space_after = Pt(6)

            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F1F5F9"/>')
            p_code._element.get_or_add_pPr().append(shd)

            r = p_code.add_run(code_text)
            r.font.name = 'Consolas'
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(30, 41, 59)
            continue

        # Table: starts with |
        if line.startswith('|') and '|' in line[1:]:
            tbl_lines = []
            while i < n and lines[i].strip().startswith('|'):
                tbl_lines.append(lines[i].strip())
                i += 1

            rows_data = []
            for tl in tbl_lines:
                if re.match(r'^\|(?:\s*:?-+:?\s*\|)+$', tl):
                    continue
                cells = [c.strip() for c in tl.split('|')[1:-1]]
                if cells:
                    rows_data.append(cells)

            if rows_data:
                num_cols = max(len(r) for r in rows_data)
                table = doc.add_table(rows=len(rows_data), cols=num_cols)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.autofit = True

                for r_idx, row_cells in enumerate(rows_data):
                    is_header = (r_idx == 0)
                    for c_idx in range(num_cols):
                        cell_text = row_cells[c_idx] if c_idx < len(row_cells) else ""
                        cell = table.cell(r_idx, c_idx)
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        _set_cell_margins(cell, top=120, bottom=120, left=150, right=150)

                        if is_header:
                            _set_cell_background(cell, "F8FAFC")

                        p_cell = cell.paragraphs[0]
                        p_cell.paragraph_format.space_before = Pt(0)
                        p_cell.paragraph_format.space_after = Pt(0)
                        _add_styled_runs(p_cell, cell_text)
                        if is_header:
                            for r in p_cell.runs:
                                r.bold = True
                                r.font.color.rgb = RGBColor(15, 23, 42)

                p_after_tbl = doc.add_paragraph()
                p_after_tbl.paragraph_format.space_after = Pt(8)
            continue

        # Headings: # H1, ## H2, ### H3, #### H4
        if line.startswith('#'):
            h_match = re.match(r'^(#{1,6})\s+(.*)$', line)
            if h_match:
                level = len(h_match.group(1))
                h_text = h_match.group(2).strip()

                p_h = doc.add_paragraph()
                p_h.paragraph_format.space_before = Pt(14 if level <= 2 else 10)
                p_h.paragraph_format.space_after = Pt(4)

                r_h = p_h.add_run(h_text)
                r_h.font.name = 'Segoe UI Semibold'
                r_h.bold = True
                if level == 1:
                    r_h.font.size = Pt(17)
                    r_h.font.color.rgb = RGBColor(15, 23, 42)
                elif level == 2:
                    r_h.font.size = Pt(14)
                    r_h.font.color.rgb = RGBColor(30, 41, 59)
                elif level == 3:
                    r_h.font.size = Pt(12)
                    r_h.font.color.rgb = RGBColor(51, 65, 85)
                else:
                    r_h.font.size = Pt(11)
                    r_h.font.color.rgb = RGBColor(71, 85, 105)

                i += 1
                continue

        # Checklist: - [ ] or - [x]
        check_match = re.match(r'^[-*]\s+\[([ xX])\]\s+(.*)$', line)
        if check_match:
            is_checked = check_match.group(1).lower() == 'x'
            item_text = check_match.group(2).strip()
            p_chk = doc.add_paragraph()
            p_chk.paragraph_format.left_indent = Inches(0.25)
            p_chk.paragraph_format.space_before = Pt(2)
            p_chk.paragraph_format.space_after = Pt(2)

            symbol = "☑ " if is_checked else "☐ "
            r_sym = p_chk.add_run(symbol)
            r_sym.font.name = 'Segoe UI Symbol'
            r_sym.bold = True
            r_sym.font.color.rgb = RGBColor(16, 185, 129) if is_checked else RGBColor(100, 116, 139)

            _add_styled_runs(p_chk, item_text)
            if is_checked:
                for r in p_chk.runs[1:]:
                    r.font.strike = True
                    r.font.color.rgb = RGBColor(148, 163, 184)

            i += 1
            continue

        # Blockquote: > text
        if line.startswith('>'):
            quote_text = re.sub(r'^>\s*', '', line)
            p_q = doc.add_paragraph()
            p_q.paragraph_format.left_indent = Inches(0.3)
            p_q.paragraph_format.space_before = Pt(4)
            p_q.paragraph_format.space_after = Pt(4)

            pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="18" w:space="10" w:color="94A3B8"/></w:pBdr>')
            p_q._element.get_or_add_pPr().append(pBdr)

            _add_styled_runs(p_q, quote_text)
            for r in p_q.runs:
                r.italic = True
                r.font.color.rgb = RGBColor(71, 85, 105)

            i += 1
            continue

        # Bullet List: - item or * item
        bullet_match = re.match(r'^[-*]\s+(.*)$', line)
        if bullet_match:
            item_text = bullet_match.group(1).strip()
            p_b = doc.add_paragraph()
            p_b.paragraph_format.left_indent = Inches(0.25)
            p_b.paragraph_format.space_before = Pt(2)
            p_b.paragraph_format.space_after = Pt(2)

            r_bullet = p_b.add_run("• ")
            r_bullet.bold = True
            r_bullet.font.color.rgb = RGBColor(100, 116, 139)

            _add_styled_runs(p_b, item_text)
            i += 1
            continue

        # Numbered List: 1. item
        num_match = re.match(r'^(\d+)\.\s+(.*)$', line)
        if num_match:
            num_str = num_match.group(1)
            item_text = num_match.group(2).strip()
            p_num = doc.add_paragraph()
            p_num.paragraph_format.left_indent = Inches(0.25)
            p_num.paragraph_format.space_before = Pt(2)
            p_num.paragraph_format.space_after = Pt(2)

            r_n = p_num.add_run(f"{num_str}. ")
            r_n.bold = True
            r_n.font.color.rgb = RGBColor(100, 116, 139)

            _add_styled_runs(p_num, item_text)
            i += 1
            continue

        # 4. Inline Drawing Directive: ::draw[id]{...} or ::draw\[id\]{...}
        draw_match = re.search(r'\\?::draw\\?\[([0-9a-zA-Z_-]+)\\?\](?:\s*\\?\{([^}]*)\\?\})?', line)
        if draw_match:
            # Check if there is text before the draw directive
            prefix_text = line[:draw_match.start()].strip()
            if prefix_text:
                p_pre = doc.add_paragraph()
                p_pre.paragraph_format.space_before = Pt(3)
                p_pre.paragraph_format.space_after = Pt(3)
                _add_styled_runs(p_pre, prefix_text)

            draw_id = draw_match.group(1)
            attr_raw = draw_match.group(2) or ""

            # Extract title attribute if present
            title_m = re.search(r'title\s*=\s*(?:"([^"]*)"|\'([^\']*)\')', attr_raw)
            draw_title = title_m.group(1) or title_m.group(2) if title_m else f"Gambar {draw_id}"

            resolved_svg = None
            resolved_png = None
            if drawing_resolver:
                try:
                    # Pass both draw_id and title to resolver
                    d_info = drawing_resolver(draw_id, title=draw_title)
                    if d_info:
                        draw_title = d_info.get("title") or draw_title
                        resolved_svg = d_info.get("svg")
                        resolved_png = d_info.get("png")
                except TypeError:
                    # Fallback if resolver only accepts 1 argument
                    try:
                        d_info = drawing_resolver(draw_id)
                        if d_info:
                            draw_title = d_info.get("title") or draw_title
                            resolved_svg = d_info.get("svg")
                            resolved_png = d_info.get("png")
                    except Exception:
                        pass
                except Exception:
                    pass

            if resolved_png:
                png_bytes = _resolve_image_bytes(resolved_png) if isinstance(resolved_png, str) else resolved_png
                if png_bytes:
                    _add_raster_image_to_doc(doc, png_bytes, alt_text=f"🎨 {draw_title}")
                elif resolved_svg and resolved_svg.strip().startswith("<svg"):
                    _add_svg_to_doc(doc, resolved_svg, title=draw_title)
            elif resolved_svg and resolved_svg.strip().startswith("<svg"):
                _add_svg_to_doc(doc, resolved_svg, title=draw_title)
            else:
                p_draw = doc.add_paragraph()
                p_draw.paragraph_format.space_before = Pt(8)
                p_draw.paragraph_format.space_after = Pt(8)
                p_draw.paragraph_format.left_indent = Inches(0.2)
                p_draw.paragraph_format.right_indent = Inches(0.2)

                shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F8FAFC"/>')
                p_draw._element.get_or_add_pPr().append(shd)

                r_d = p_draw.add_run(f"🎨 [Gambar/Canvas: {draw_title}]")
                r_d.bold = True
                r_d.font.color.rgb = RGBColor(100, 116, 139)

            suffix_text = line[draw_match.end():].strip()
            if suffix_text:
                p_suf = doc.add_paragraph()
                p_suf.paragraph_format.space_before = Pt(3)
                p_suf.paragraph_format.space_after = Pt(3)
                _add_styled_runs(p_suf, suffix_text)

            i += 1
            continue

        # 5. Standalone Image (matches ![alt](url), !image.png, ![image.png], [image.png], <img...>, etc.)
        img_info = _parse_standalone_image(line)
        if img_info:
            img_src, alt_text = img_info
            img_bytes = _resolve_image_bytes(img_src, image_resolver, alt=alt_text)
            if img_bytes:
                _add_raster_image_to_doc(doc, img_bytes, alt_text=alt_text)
            else:
                p_img = doc.add_paragraph()
                r_img = p_img.add_run(f"🖼️ [{alt_text or 'Gambar'}]")
                r_img.font.color.rgb = RGBColor(100, 116, 139)
            i += 1
            continue

        # 7. Standard Paragraph (with potential inline images or styled runs)
        if '![' in line and '](' in line:
            segments = re.split(r'(!\[[^\]]*\]\([^)]+\))', line)
            for seg in segments:
                if not seg:
                    continue
                sub_img_m = re.match(r'^!\[([^\]]*)\]\(([^)]+)\)$', seg)
                if sub_img_m:
                    alt_text = sub_img_m.group(1)
                    img_src = sub_img_m.group(2)
                    img_bytes = _resolve_image_bytes(img_src, image_resolver)
                    if img_bytes:
                        _add_raster_image_to_doc(doc, img_bytes, alt_text=alt_text)
                    else:
                        p_sub = doc.add_paragraph()
                        r_sub = p_sub.add_run(f"🖼️ [{alt_text or 'Gambar'}]")
                        r_sub.font.color.rgb = RGBColor(100, 116, 139)
                else:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(3)
                    p.paragraph_format.space_after = Pt(5)
                    _add_styled_runs(p, seg)
            i += 1
            continue

        # Regular Paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(5)
        _add_styled_runs(p, line)
        i += 1

    out_io = io.BytesIO()
    doc.save(out_io)
    out_io.seek(0)
    return out_io
