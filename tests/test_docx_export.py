import io
import docx
import pytest
from starlette.testclient import TestClient
import webapp
from docx_exporter import markdown_to_docx


def test_markdown_to_docx_structure():
    title = "Laporan Proyek TaskFlow"
    content = """# Pendahuluan

Berikut adalah ringkasan progres:
- [ ] Task 1: Desain UI
- [x] Task 2: Implementasi API

## Rincian Data
Teks **tebal**, *miring*, `kode inline`, dan [link](https://yatno.web.id).

> Catatan penting: Pastikan semua data tervalidasi.

| Modul | Status | PIC |
| :--- | :---: | ---: |
| Auth | Selesai | Yatno |
| Notes | Testing | Tim Dev |

```python
def halo():
    return "Hello TaskFlow"
```
"""
    meta = {
        "tags": ["proyek", "laporan"],
        "updated_at": "2026-08-21T09:00:00"
    }

    doc_io = markdown_to_docx(title, content, meta)
    assert isinstance(doc_io, io.BytesIO)
    doc_io.seek(0)

    # Verify DOCX is valid and can be loaded by python-docx
    doc = docx.Document(doc_io)
    paragraphs = [p.text for p in doc.paragraphs]
    
    # Check title and metadata presence
    assert any("Laporan Proyek TaskFlow" in p for p in paragraphs)
    assert any("#proyek" in p for p in paragraphs)
    
    # Check headings
    assert any("Pendahuluan" in p for p in paragraphs)
    assert any("Rincian Data" in p for p in paragraphs)
    
    # Check checklist symbols
    assert any("☐ Task 1: Desain UI" in p for p in paragraphs)
    assert any("☑ Task 2: Implementasi API" in p for p in paragraphs)
    
    # Check table
    assert len(doc.tables) >= 1
    table = doc.tables[0]
    headers = [c.text.strip() for c in table.rows[0].cells]
    assert "Modul" in headers
    assert "Status" in headers
    assert "PIC" in headers
    row1 = [c.text.strip() for c in table.rows[1].cells]
    assert "Auth" in row1
    assert "Selesai" in row1


def test_export_endpoints():
    client = TestClient(webapp.app)
    
    # Register / login a test user
    import uuid
    uname = f"expuser_{uuid.uuid4().hex[:6]}"
    reg = client.post("/api/auth/register", json={"username": uname, "email": f"{uname}@test.id", "password": "Password123!"})
    token = reg.json().get("token") or reg.json().get("access_token")
    headers = {"Authorization": f"Bearer {token}"} if token else {}

    # Create note
    create_res = client.post("/api/scratchpad", json={
        "title": "Dokumen Uji Export",
        "content": "# Halo Dunia\n\n- [ ] Item 1\n- [x] Item 2\n\n| Kolom A | Kolom B |\n|---|---|\n| 100 | 200 |",
        "tags": ["test", "export"]
    }, headers=headers)
    assert create_res.status_code == 200
    note_id = create_res.json()["id"]

    # Test GET /api/scratchpad/{id}/export/docx
    docx_res = client.get(f"/api/scratchpad/{note_id}/export/docx", headers=headers)
    assert docx_res.status_code == 200
    assert "wordprocessingml.document" in docx_res.headers.get("content-type", "")
    assert f'filename="Dokumen Uji Export.docx"' in docx_res.headers.get("content-disposition", "")
    
    # Verify downloaded docx content
    doc = docx.Document(io.BytesIO(docx_res.content))
    assert any("Dokumen Uji Export" in p.text for p in doc.paragraphs)
    assert len(doc.tables) >= 1

    # Test GET /api/scratchpad/{id}/export/md
    md_res = client.get(f"/api/scratchpad/{note_id}/export/md", headers=headers)
    assert md_res.status_code == 200
    assert "text/markdown" in md_res.headers.get("content-type", "")
    assert f'filename="Dokumen Uji Export.md"' in md_res.headers.get("content-disposition", "")
    assert "# Halo Dunia" in md_res.text

    # Test POST /api/scratchpad/export/docx (live payload)
    post_res = client.post("/api/scratchpad/export/docx", json={
        "title": "Catatan Langsung",
        "content": "Ini konten langsung dari editor."
    }, headers=headers)
    assert post_res.status_code == 200
    assert "wordprocessingml.document" in post_res.headers.get("content-type", "")
    assert 'filename="Catatan Langsung.docx"' in post_res.headers.get("content-disposition", "")


def test_markdown_to_docx_with_images_and_drawings():
    import base64
    from PIL import Image

    # Create dummy PNG bytes
    img = Image.new('RGB', (100, 50), color=(16, 185, 129))
    buf = io.BytesIO()
    img.save(buf, format='PNG')
    png_bytes = buf.getvalue()
    b64_str = f"data:image/png;base64,{base64.b64encode(png_bytes).decode('ascii')}"

    svg_preview_str = '<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 400 200"><rect width="400" height="200" fill="#fef3c7"/><text x="200" y="100">Diagram</text></svg>'

    title = "Catatan dengan Gambar & Canvas"
    content = f"""# Gambar dan Diagram

Berikut adalah gambar base64:
![Gambar Dummy]({b64_str})

Berikut adalah standalone bang image:
!image.png

Berikut adalah standalone bracket image:
![screenshot.jpg]

Berikut adalah HTML img:
<img src="/api/scratchpad/attachments/88/view" alt="Foto HTML" />

Berikut adalah drawing inline:
::draw[101]{{size="M" title="Diagram Alur"}}

Berikut adalah lampiran gambar resolver:
![Foto Lampiran](/api/scratchpad/attachments/55/view)
"""

    def mock_drawing_resolver(did, title=None):
        if str(did) == "101":
            return {
                "title": "Diagram Alur",
                "svg": svg_preview_str
            }
        return None

    def mock_image_resolver(src, alt=None):
        if "55" in src or "88" in src or "image.png" in src or "screenshot.jpg" in src:
            return png_bytes
        return None

    doc_io = markdown_to_docx(
        title,
        content,
        drawing_resolver=mock_drawing_resolver,
        image_resolver=mock_image_resolver
    )
    assert isinstance(doc_io, io.BytesIO)
    doc = docx.Document(doc_io)

    # Verify paragraphs and captions
    paragraphs = [p.text for p in doc.paragraphs]
    assert any("Gambar dan Diagram" in p for p in paragraphs)
    assert any("Diagram Alur" in p for p in paragraphs)
    assert any("Foto Lampiran" in p for p in paragraphs)
    assert any("Foto HTML" in p for p in paragraphs)

    # Verify images exist in docx package
    image_parts = [p for p in doc.part.package.parts if 'image' in p.partname]
    assert len(image_parts) >= 1


